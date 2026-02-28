import fitz
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import zipfile
import re


class PDFService:

    # ---------- Helpers ----------

    @staticmethod
    def limpiar_parrafo(texto):
        return texto.replace("\n", " ").strip()

    @staticmethod
    def es_titulo(texto):
        return len(texto) < 60 and (texto.isupper() or texto.istitle())

    @staticmethod
    def es_subtitulo(texto):
        return len(texto) < 90 and texto.istitle()

    @staticmethod
    def es_link(texto):
        return texto.startswith("http")

    @staticmethod
    def es_tabla_real(lineas):
        tabla_actual = []
        for l in lineas:
            cols = [c.strip() for c in re.split(r"\t|  ", l) if c.strip()]
            if len(cols) >= 2:
                tabla_actual.append(l)
            else:
                if len(tabla_actual) >= 2:
                    return tabla_actual
                tabla_actual = []
        if len(tabla_actual) >= 2:
            return tabla_actual
        return None

    @staticmethod
    def construir_tabla_limpia(word, lineas):
        max_cols = max(len(re.split(r"\t|  ", l.strip())) for l in lineas)
        tabla = word.add_table(rows=len(lineas), cols=max_cols)
        tabla.style = "Table Grid"

        for i, linea in enumerate(lineas):
            cols = [c.strip() for c in re.split(r"\t|  ", linea) if c.strip()]
            for j, col in enumerate(cols):
                cell = tabla.rows[i].cells[j]
                if PDFService.es_subtitulo(col):
                    p = cell.paragraphs[0]
                    run = p.add_run(col)
                    run.font.size = Pt(11)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    cell.text = col

    @staticmethod
    def procesar_pagina(word, page):
        bloques = page.get_text("dict")["blocks"]

        for block in bloques:
            if "lines" not in block:
                continue

            lineas_texto = [
                PDFService.limpiar_parrafo(
                    " ".join(span["text"] for span in linea["spans"])
                )
                for linea in block["lines"]
            ]

            tabla_detectada = PDFService.es_tabla_real(lineas_texto)
            if tabla_detectada:
                PDFService.construir_tabla_limpia(word, tabla_detectada)
                continue

            for line in block["lines"]:
                linea_texto = " ".join(span["text"] for span in line["spans"]).strip()
                if not linea_texto:
                    continue

                if PDFService.es_link(linea_texto):
                    p = word.add_paragraph()
                    run = p.add_run(linea_texto)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
                    continue

                if PDFService.es_titulo(linea_texto):
                    word.add_heading(linea_texto, level=1)

                elif PDFService.es_subtitulo(linea_texto):
                    word.add_heading(linea_texto, level=2)

                else:
                    p = word.add_paragraph()
                    run = p.add_run(linea_texto)
                    run.font.size = Pt(11)

    # ---------- Conversiones ----------

    @staticmethod
    def pdf_to_png(pdf_bytes):
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for i, page in enumerate(doc):
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_name = f"page_{i+1}.png"
                zipf.writestr(img_name, pix.tobytes("png"))

        zip_buffer.seek(0)
        doc.close()
        return zip_buffer

    @staticmethod
    def pdf_to_word(pdf_bytes):
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        word = Document()

        style = word.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for i, page in enumerate(doc):
            PDFService.procesar_pagina(word, page)
            if i < len(doc) - 1:
                word.add_page_break()

        buffer = BytesIO()
        word.save(buffer)
        buffer.seek(0)
        doc.close()

        return buffer