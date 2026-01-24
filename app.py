from flask import Flask, request, send_file, render_template
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import zipfile
from io import BytesIO
import re
from flask import jsonify
import smtplib
from email.message import EmailMessage



app = Flask(__name__)

OUTPUT_FOLDER = "static/outputs"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ---------------- Funciones auxiliares ----------------

def limpiar_texto(texto):
    return " ".join(texto.split())

def limpiar_parrafo(texto):
    return texto.replace("\n", " ").strip()

def es_titulo(texto):
    return len(texto) < 60 and (texto.isupper() or texto.istitle())

def es_subtitulo(texto):
    return len(texto) < 90 and texto.istitle()

def es_link(texto):
    return texto.startswith("http")

def agrupar_por_filas(block):
    filas = {}
    for line in block["lines"]:
        y = round(line["bbox"][1], 1)
        filas.setdefault(y, []).extend(line["spans"])
    return filas

def construir_tabla_limpia(word, lineas):
    """
    Construye la tabla y formatea subtítulos dentro de la tabla.
    """
    max_cols = max(len(re.split(r"\t|  ", l.strip())) for l in lineas)
    tabla = word.add_table(rows=len(lineas), cols=max_cols)
    tabla.style = "Table Grid"

    for i, linea in enumerate(lineas):
        cols = [c.strip() for c in re.split(r"\t|  ", linea) if c.strip()]
        for j, col in enumerate(cols):
            cell = tabla.rows[i].cells[j]
            if es_subtitulo(col):
                p = cell.paragraphs[0]
                run = p.add_run(col)
                run.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell.text = col

def es_tabla_real(lineas):
    """
    Detecta tablas reales: al menos 2 líneas consecutivas con 2+ columnas
    """
    tabla_actual = []
    for l in lineas:
        cols = [c.strip() for c in re.split(r"\t|  ", l) if c.strip()]
        if len(cols) >= 2:
            tabla_actual.append(l)
        else:
            if len(tabla_actual) >= 2:
                return tabla_actual
            tabla_actual = []
    if len(tabla_actual) >= 2:
        return tabla_actual
    return None

# ---------------- Procesamiento de página ----------------

def procesar_pagina(word, page):
    bloques = page.get_text("dict")["blocks"]

    for block in bloques:
        if "lines" not in block:
            continue

        filas = agrupar_por_filas(block)

        lineas_texto = [limpiar_parrafo(" ".join(span["text"] for span in linea["spans"]))
                        for linea in block["lines"]]

        # Detectar tablas reales
        tabla_detectada = es_tabla_real(lineas_texto)
        if tabla_detectada:
            construir_tabla_limpia(word, tabla_detectada)
            continue

        # Procesar línea por línea
        for line in block["lines"]:
            linea_texto = " ".join(span["text"] for span in line["spans"]).strip()
            if not linea_texto:
                continue

            # Links
            if es_link(linea_texto):
                p = word.add_paragraph()
                run = p.add_run(linea_texto)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                continue

            # Títulos
            if es_titulo(linea_texto):
                p = word.add_heading(linea_texto, level=1)
                x0 = line["spans"][0]["bbox"][0]
                x1 = line["spans"][-1]["bbox"][2]
                page_width = page.rect.width
                if (x0 + x1)/2 > page_width*0.4 and (x0 + x1)/2 < page_width*0.6:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Subtítulos
            elif es_subtitulo(linea_texto):
                p = word.add_heading(linea_texto, level=2)
                x0 = line["spans"][0]["bbox"][0]
                x1 = line["spans"][-1]["bbox"][2]
                page_width = page.rect.width
                if (x0 + x1)/2 > page_width*0.4 and (x0 + x1)/2 < page_width*0.6:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Párrafos normales
            else:
                p = word.add_paragraph()
                run = p.add_run(linea_texto)
                run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(6)

# ---------------- Rutas Flask ----------------

@app.route("/feedback", methods=["POST"])
def feedback():
    data = request.get_json()
    print("Feedback recibido:", data.get("feedback"))
    return jsonify({"status": "ok"})


@app.route("/gracias")
def gracias():
    return render_template("gracias.html")


@app.route("/terms")
def terms():
    return render_template("terms.html")



@app.route("/contact")
def contact():
    return render_template("contact.html")



@app.route("/")
def home():
    return render_template("index.html")

@app.route("/privacy")
def privacy():
    return render_template("privacy.html")





# PDF → PNG
@app.route("/convert/png", methods=["POST"])
def pdf_to_png():
    try:
        if "file" not in request.files:
            return "No se envió ningún archivo", 400

        pdf_file = request.files["file"]
        if pdf_file.filename == "":
            return "Archivo vacío", 400

        pdf_bytes = pdf_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for i, page in enumerate(doc):
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 🔥 FIX
                img_name = f"page_{i+1}.png"
                img_bytes = pix.tobytes("png")
                zipf.writestr(img_name, img_bytes)

        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name="imagenes.zip",
            mimetype="application/zip"
        )

    except Exception as e:
        print("ERROR PNG:", e)
        return f"Error procesando el PDF: {str(e)}", 500

# PDF → WORD
@app.route("/convert/word", methods=["POST"])
def pdf_to_word():
    try:
        if "file" not in request.files:
            return "No se envió archivo", 400

        pdf_file = request.files["file"]
        pdf_bytes = pdf_file.read()

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        word = Document()

        style = word.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        for i, page in enumerate(doc):
            procesar_pagina(word, page)
            if i < len(doc) - 1:
                word.add_page_break()

        buffer = BytesIO()
        word.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="convertido.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print("ERROR WORD:", e)
        return f"Error procesando PDF a Word: {e}", 500



if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)



